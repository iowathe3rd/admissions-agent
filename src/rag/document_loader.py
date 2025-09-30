"""
Модуль для загрузки различных форматов документов с строгой типизацией.
Поддерживает: JSON, TXT, PDF, DOCX
"""

import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Protocol, Union

# Импорты с обработкой ошибок для опциональных зависимостей
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    PyPDF2 = None  # type: ignore
    PdfReader = None  # type: ignore

try:
    from docx import Document as DocxDocument
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None  # type: ignore
    docx = None  # type: ignore

logger = logging.getLogger(__name__)


class DocumentData(Protocol):
    """Протокол для данных документа."""
    source: str
    text: str


class LoaderResult:
    """Результат загрузки документа."""
    
    def __init__(self, source: str, text: str, success: bool = True, error: Optional[str] = None):
        self.source = source
        self.text = text
        self.success = success
        self.error = error
    
    def to_dict(self) -> Dict[str, str]:
        """Преобразует в словарь для совместимости."""
        return {"source": self.source, "text": self.text}


class DocumentLoader:
    """Загрузчик документов различных форматов с строгой типизацией."""
    
    def __init__(self, encoding: str = 'utf-8'):
        self.encoding = encoding
        self.supported_extensions = {'.json', '.txt', '.pdf', '.docx'}
    
    def load_json_file(self, file_path: Path) -> List[LoaderResult]:
        """Загружает JSON файл и возвращает список документов."""
        results: List[LoaderResult] = []
        
        try:
            logger.info(f"Обрабатываем JSON файл: {file_path.name}")
            
            with open(file_path, 'r', encoding=self.encoding) as f:
                data = json.load(f)
            
            if not isinstance(data, list):
                return [LoaderResult(
                    source=file_path.stem,
                    text="",
                    success=False,
                    error="JSON должен содержать массив объектов"
                )]
            
            source_name = file_path.stem
            
            for i, item in enumerate(data):
                if not isinstance(item, dict):
                    logger.warning(f"Элемент {i} в {file_path.name} не является объектом")
                    continue
                
                # Безопасное преобразование значений в строки
                text_parts = []
                for key, value in item.items():
                    if value is not None:
                        # Обработка списков и вложенных объектов
                        if isinstance(value, (list, tuple)):
                            value_str = ", ".join(str(v) for v in value)
                        elif isinstance(value, dict):
                            value_str = "; ".join(f"{k}: {v}" for k, v in value.items())
                        else:
                            value_str = str(value)
                        
                        text_parts.append(f"{key}: {value_str}")
                
                text = " ".join(text_parts)
                if text.strip():  # Только непустые документы
                    results.append(LoaderResult(source=source_name, text=text))
            
            logger.info(f"Загружено {len(results)} документов из {file_path.name}")
            
        except json.JSONDecodeError as e:
            error_msg = f"Ошибка парсинга JSON в {file_path.name}: {e}"
            logger.error(error_msg)
            return [LoaderResult(source=file_path.stem, text="", success=False, error=error_msg)]
        
        except Exception as e:
            error_msg = f"Ошибка при загрузке {file_path.name}: {e}"
            logger.error(error_msg)
            return [LoaderResult(source=file_path.stem, text="", success=False, error=error_msg)]
        
        return results
    
    def load_txt_file(self, file_path: Path) -> List[LoaderResult]:
        """Загружает TXT файл."""
        try:
            logger.info(f"Обрабатываем TXT файл: {file_path.name}")
            
            with open(file_path, 'r', encoding=self.encoding) as f:
                text = f.read().strip()
            
            if not text:
                return [LoaderResult(
                    source=file_path.stem,
                    text="",
                    success=False,
                    error="Файл пустой"
                )]
            
            # Нормализация текста
            text = re.sub(r'\s+', ' ', text)  # Убираем лишние пробелы
            
            return [LoaderResult(source=file_path.stem, text=text)]
            
        except UnicodeDecodeError:
            # Пробуем другие кодировки
            for encoding in ['cp1251', 'latin1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read().strip()
                    logger.info(f"Файл {file_path.name} загружен с кодировкой {encoding}")
                    return [LoaderResult(source=file_path.stem, text=text)]
                except UnicodeDecodeError:
                    continue
            
            error_msg = f"Не удалось определить кодировку файла {file_path.name}"
            logger.error(error_msg)
            return [LoaderResult(source=file_path.stem, text="", success=False, error=error_msg)]
        
        except Exception as e:
            error_msg = f"Ошибка при загрузке TXT файла {file_path.name}: {e}"
            logger.error(error_msg)
            return [LoaderResult(source=file_path.stem, text="", success=False, error=error_msg)]
    
    def load_pdf_file(self, file_path: Path) -> List[LoaderResult]:
        """Загружает PDF файл."""
        if not PDF_AVAILABLE or PdfReader is None:
            error_msg = "PyPDF2 не установлен. Установите: pip install PyPDF2"
            logger.error(error_msg)
            return [LoaderResult(source=file_path.stem, text="", success=False, error=error_msg)]
        
        try:
            logger.info(f"Обрабатываем PDF файл: {file_path.name}")
            
            text_parts: List[str] = []
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)  # type: ignore
                
                if len(pdf_reader.pages) == 0:
                    return [LoaderResult(
                        source=file_path.stem,
                        text="",
                        success=False,
                        error="PDF файл не содержит страниц"
                    )]
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text.strip())
                    except Exception as e:
                        logger.warning(f"Ошибка извлечения текста со страницы {page_num + 1}: {e}")
                        continue
            
            if not text_parts:
                return [LoaderResult(
                    source=file_path.stem,
                    text="",
                    success=False,
                    error="Не удалось извлечь текст из PDF"
                )]
            
            # Объединяем текст со страниц
            full_text = "\n".join(text_parts)
            
            # Нормализация текста
            full_text = re.sub(r'\s+', ' ', full_text)
            full_text = full_text.strip()
            
            return [LoaderResult(source=file_path.stem, text=full_text)]
            
        except Exception as e:
            error_msg = f"Ошибка при загрузке PDF файла {file_path.name}: {e}"
            logger.error(error_msg)
            return [LoaderResult(source=file_path.stem, text="", success=False, error=error_msg)]
    
    def load_docx_file(self, file_path: Path) -> List[LoaderResult]:
        """Загружает DOCX файл."""
        if not DOCX_AVAILABLE or docx is None:
            error_msg = "python-docx не установлен. Установите: pip install python-docx"
            logger.error(error_msg)
            return [LoaderResult(source=file_path.stem, text="", success=False, error=error_msg)]
        
        try:
            logger.info(f"Обрабатываем DOCX файл: {file_path.name}")
            
            doc = docx.Document(file_path)  # type: ignore
            
            text_parts: List[str] = []
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    text_parts.append(paragraph_text)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))
            
            if not text_parts:
                return [LoaderResult(
                    source=file_path.stem,
                    text="",
                    success=False,
                    error="DOCX файл не содержит текста"
                )]
            
            # Объединяем весь текст
            full_text = "\n".join(text_parts)
            
            # Нормализация текста
            full_text = re.sub(r'\s+', ' ', full_text)
            full_text = full_text.strip()
            
            return [LoaderResult(source=file_path.stem, text=full_text)]
            
        except Exception as e:
            error_msg = f"Ошибка при загрузке DOCX файла {file_path.name}: {e}"
            logger.error(error_msg)
            return [LoaderResult(source=file_path.stem, text="", success=False, error=error_msg)]
    
    def load_file(self, file_path: Path) -> List[LoaderResult]:
        """Загружает файл на основе его расширения."""
        suffix = file_path.suffix.lower()
        
        if suffix not in self.supported_extensions:
            error_msg = f"Неподдерживаемый формат файла: {suffix}"
            logger.warning(error_msg)
            return [LoaderResult(source=file_path.stem, text="", success=False, error=error_msg)]
        
        if suffix == '.json':
            return self.load_json_file(file_path)
        elif suffix == '.txt':
            return self.load_txt_file(file_path)
        elif suffix == '.pdf':
            return self.load_pdf_file(file_path)
        elif suffix == '.docx':
            return self.load_docx_file(file_path)
        else:
            error_msg = f"Обработчик для {suffix} не реализован"
            logger.error(error_msg)
            return [LoaderResult(source=file_path.stem, text="", success=False, error=error_msg)]
    
    def load_directory(self, directory_path: Path) -> List[LoaderResult]:
        """Загружает все поддерживаемые файлы из директории."""
        all_results: List[LoaderResult] = []
        
        if not directory_path.exists():
            logger.error(f"Директория не найдена: {directory_path}")
            return []
        
        if not directory_path.is_dir():
            logger.error(f"Путь не является директорией: {directory_path}")
            return []
        
        # Собираем все поддерживаемые файлы
        supported_files: List[Path] = []
        for extension in self.supported_extensions:
            pattern = f"*{extension}"
            supported_files.extend(directory_path.glob(pattern))
        
        if not supported_files:
            logger.warning(f"Не найдено поддерживаемых файлов в {directory_path}")
            return []
        
        logger.info(f"Найдено {len(supported_files)} файлов для обработки")
        
        # Обрабатываем каждый файл
        successful_files = 0
        failed_files = 0
        
        for file_path in sorted(supported_files):
            file_results = self.load_file(file_path)
            
            for result in file_results:
                all_results.append(result)
                if result.success:
                    successful_files += 1
                else:
                    failed_files += 1
        
        logger.info(f"Обработка завершена: {successful_files} успешно, {failed_files} с ошибками")
        
        return all_results
    
    def get_statistics(self, results: List[LoaderResult]) -> Dict[str, Any]:
        """Возвращает статистику загрузки."""
        successful = [r for r in results if r.success]
        failed = [r for r in results if not r.success]
        
        total_chars = sum(len(r.text) for r in successful)
        avg_length = total_chars / len(successful) if successful else 0
        
        return {
            "total_documents": len(results),
            "successful": len(successful),
            "failed": len(failed),
            "total_characters": total_chars,
            "average_length": avg_length,
            "errors": [{"source": r.source, "error": r.error} for r in failed if r.error]
        }